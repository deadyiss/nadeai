from pathlib import Path
from typing import Optional

import pypdf
from docx import Document as DocxDocument

from core.ocr_engine import get_ocr_engine
from utils.file_validator import validate_file
from utils.logger import get_logger
from utils.text_cleaner import clean_text, count_words
from utils.date_extractor import extract_dates

logger = get_logger(__name__)


class DocumentProcessor:
    def __init__(self):
        self._ocr = None  # lazy-load

    def _get_ocr(self):
        if self._ocr is None:
            self._ocr = get_ocr_engine()
        return self._ocr

    def _extract_pdf(self, filepath: str) -> str:
        logger.info(f"Extracting PDF: {filepath}")
        text_parts = []
        with open(filepath, "rb") as f:
            reader = pypdf.PdfReader(f)
            for i, page in enumerate(reader.pages):
                try:
                    txt = page.extract_text() or ""
                    if txt.strip():
                        text_parts.append(txt)
                except Exception as e:
                    logger.warning(f"Failed extract page {i}: {e}")
        full_text = "\n\n".join(text_parts)

        # If PDF has very little text -> probably scanned, fallback to OCR
        if len(full_text.strip()) < 50:
            logger.info(f"PDF text too short ({len(full_text.strip())} chars), falling back to OCR")
            try:
                ocr_text, _ = self._extract_image(filepath)
                if ocr_text and ocr_text.strip():
                    return ocr_text
            except Exception as e:
                logger.warning(f"OCR fallback failed for PDF: {e}")
        return full_text

    def _extract_docx(self, filepath: str) -> str:
        logger.info(f"Extracting DOCX: {filepath}")
        doc = DocxDocument(filepath)
        parts = []

        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)

        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    parts.append(row_text)

        return "\n".join(parts)

    def _extract_image(self, filepath: str) -> tuple[str, float]:
        logger.info(f"Extracting image via OCR: {filepath}")
        ocr = self._get_ocr()
        result = ocr.extract(filepath)
        return result["text"], result["confidence"]

    def process(self, filepath: str, original_filename: Optional[str] = None) -> dict:
        """
        Process a file and return:
        {
            "filename": str,
            "text": str (cleaned),
            "word_count": int,
            "dates": list[str] (ISO format),
            "extension": str,
            "ocr_confidence": float | None,
        }
        """
        info = validate_file(filepath, original_filename)
        ext = info["extension"]
        ocr_conf: Optional[float] = None

        if info["is_pdf"]:
            raw_text = self._extract_pdf(filepath)
        elif info["is_docx"]:
            raw_text = self._extract_docx(filepath)
        elif info["is_txt"]:
            with open(filepath, "r", encoding="utf-8", errors="ignore") as f:
                raw_text = f.read()
        elif info["is_image"]:
            raw_text, ocr_conf = self._extract_image(filepath)
        else:
            raise ValueError(f"Unsupported extension: {ext}")

        cleaned = clean_text(raw_text)
        if not cleaned:
            raise ValueError("No text could be extracted from file")

        return {
            "filename": info["filename"],
            "text": cleaned,
            "word_count": count_words(cleaned),
            "dates": extract_dates(cleaned),
            "extension": ext,
            "ocr_confidence": ocr_conf,
        }


_processor: Optional[DocumentProcessor] = None


def get_document_processor() -> DocumentProcessor:
    global _processor
    if _processor is None:
        _processor = DocumentProcessor()
    return _processor